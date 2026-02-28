"""
File Reader Utility for Resume Screening AI
Handles PDF and DOCX file reading.
"""

import os


def read_pdf(file_path):
    """
    Read text from a PDF file.
    """
    try:
        import PyPDF2
        
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        raise Exception(f"Error reading PDF file: {str(e)}")


def read_docx(file_path):
    """
    Read text from a DOCX file.
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also read tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {str(e)}")


def read_txt(file_path):
    """
    Read text from a plain text file.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        raise Exception(f"Error reading text file: {str(e)}")


def read_resume_file(file_path):
    """
    Read resume from various file formats.
    Supported formats: PDF, DOCX, TXT
    """
    if not os.path.exists(file_path):
        raise Exception(f"File not found: {file_path}")
    
    # Get file extension
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext == '.pdf':
        return read_pdf(file_path)
    elif ext == '.docx':
        return read_docx(file_path)
    elif ext == '.txt':
        return read_txt(file_path)
    else:
        raise Exception(f"Unsupported file format: {ext}. Supported formats: PDF, DOCX, TXT")


def read_uploaded_file(file_storage):
    """
    Read text from an uploaded file (Flask FileStorage object).
    """
    import tempfile
    import os
    
    # Get the filename
    filename = file_storage.filename
    
    # Check file extension
    _, ext = os.path.splitext(filename)
    ext = ext.lower()
    
    if ext not in ['.pdf', '.docx', '.txt']:
        raise Exception(f"Unsupported file format: {ext}. Supported formats: PDF, DOCX, TXT")
    
    # Save to temp file
    temp_path = os.path.join(tempfile.gettempdir(), filename)
    file_storage.save(temp_path)
    
    try:
        # Read the file
        text = read_resume_file(temp_path)
        return text
    finally:
        # Clean up temp file
        if os.path.exists(temp_path):
            os.remove(temp_path)
